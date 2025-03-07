from reportlab.lib.pagesizes import A4
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak,
    Image, Table, TableStyle, Frame, PageTemplate
)
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import mm
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from typing import List, Dict
import logging
from datetime import datetime
from config import PDF_CONFIG
import os

logger = logging.getLogger(__name__)

class PDFGenerator:
    def __init__(self, articles: List[Dict], filename: str):
        self.articles = articles
        self.filename = filename
        self.styles = self._create_styles()
        self.logo = self._get_logo()

    def _get_logo(self):
        if os.path.exists(PDF_CONFIG["branding"]["logo_path"]):
            return Image(PDF_CONFIG["branding"]["logo_path"], width=40*mm, height=12*mm)
        return None

    def _create_styles(self):
        cfg = PDF_CONFIG
        return {
            'h1': ParagraphStyle(
                name='Heading1',
                fontName=cfg["fonts"]["heading"],
                fontSize=cfg["fonts"]["sizes"]["h1"],
                textColor=colors.HexColor(cfg["colors"]["primary"]),
                leading=cfg["fonts"]["sizes"]["h1"] * 1.2,
                spaceAfter=cfg["spacing"]["section"],
                alignment=0
            ),
            'h2': ParagraphStyle(
                name='Heading2',
                fontName=cfg["fonts"]["heading"],
                fontSize=cfg["fonts"]["sizes"]["h2"],
                textColor=colors.HexColor(cfg["colors"]["primary"]),
                spaceAfter=cfg["spacing"]["paragraph"]
            ),
            'meta': ParagraphStyle(
                name='Meta',
                fontName=cfg["fonts"]["accent"],
                fontSize=cfg["fonts"]["sizes"]["meta"],
                textColor=colors.HexColor(cfg["colors"]["secondary"]),
                spaceAfter=cfg["spacing"]["paragraph"]
            ),
            'body': ParagraphStyle(
                name='Body',
                fontName=cfg["fonts"]["body"],
                fontSize=cfg["fonts"]["sizes"]["body"],
                textColor=colors.black,  # Texto siempre negro
                leading=cfg["fonts"]["sizes"]["body"] * cfg["spacing"]["line_height"],
                spaceAfter=cfg["spacing"]["paragraph"]
            )
        }
    def _header_footer(self, canvas, doc):
        canvas.saveState()
        # Encabezado verde
        canvas.setFillColor(colors.HexColor(PDF_CONFIG["colors"]["primary"]))
        canvas.rect(0, A4[1] - 40, A4[0], 40, fill=1, stroke=0)
        
        if self.logo:
            self.logo.drawOn(canvas, 35*mm, A4[1] - 35*mm)
        # Footer
        canvas.setFont(PDF_CONFIG["fonts"]["body"], PDF_CONFIG["fonts"]["sizes"]["meta"])
        canvas.setFillColor(colors.HexColor(PDF_CONFIG["colors"]["secondary"]))
        footer_text = f"{PDF_CONFIG['branding']['website']} - Página {doc.page}"
        canvas.drawCentredString(A4[0]/2, 15*mm, footer_text)
        canvas.restoreState()

    def generate(self):
        try:
            doc = SimpleDocTemplate(
                self.filename,
                pagesize=A4,
                leftMargin=PDF_CONFIG["page"]["margin"]["left"] * mm,
                rightMargin=PDF_CONFIG["page"]["margin"]["right"] * mm,
                topMargin=PDF_CONFIG["page"]["margin"]["top"] * mm,
                bottomMargin=PDF_CONFIG["page"]["margin"]["bottom"] * mm
            )
            
            elements = [
                Spacer(1, 20),
                Paragraph("Blog Cultivo Loco", self.styles['h1']),
                Table(
                    [[""]],
                    colWidths=["100%"],
                    style=[
                        ('LINEBELOW', (0,0), (-1,-1), 1, colors.HexColor(PDF_CONFIG["colors"]["border"]))
                    ]
                ),
                Spacer(1, PDF_CONFIG["spacing"]["section"])
            ]
            
            for article in self.articles:
                elements += [
                    Paragraph(article["title"], self.styles['h2']),
                    Paragraph(f"Publicado el {article['date']}", self.styles['meta']),
                    Spacer(1, 8),
                    Paragraph(article["content"], self.styles['body']),
                    self._create_divider(),
                    PageBreak()
                ]
            
            doc.build(elements, onFirstPage=self._header_footer, onLaterPages=self._header_footer)
            logging.info(f"PDF generado: {self.filename}")
        
        except Exception as e:
            logging.error(f"Error generando PDF: {str(e)}")
            raise

    def _create_divider(self):
        return Table(
            [[""]],
            colWidths=["100%"],
            style=[
                ('LINEABOVE', (0,0), (-1,-1), 1, colors.HexColor(PDF_CONFIG["colors"]["border"]))
            ]
        )

class DOCXGenerator:
    def __init__(self, articles: List[Dict], filename: str):
        self.articles = articles
        self.filename = filename
    
    def generate(self):
        try:
            doc = Document()
            self._setup_styles(doc)
            
            # Portada
            doc.add_heading('Catálogo de Artículos', 0)
            doc.add_paragraph().add_run('Edición Especial').italic = True
            
            # Contenido
            for article in self.articles:
                doc.add_heading(article['title'], level=1)
                doc.add_paragraph(f"Publicado el {article['date']}", style='Intense Quote')
                doc.add_paragraph(article['content'])
                doc.add_page_break()
            
            doc.save(self.filename)
            logger.info(f"DOCX generado: {self.filename}")
        
        except Exception as e:
            logger.error(f"Error generando DOCX: {str(e)}")
            raise
    
    def _setup_styles(self, doc):
        cfg = PDF_CONFIG
        styles = doc.styles
        
        # Título principal
        title_style = styles['Title']
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(cfg["fonts"]["sizes"]["h1"])
        title_style.font.color.rgb = RGBColor.from_string(cfg["colors"]["primary"][1:])
        
        # Encabezado artículo
        heading_style = styles['Heading1']
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(cfg["fonts"]["sizes"]["h2"])
        heading_style.font.color.rgb = RGBColor.from_string(cfg["colors"]["primary"][1:])
        
        # Fecha
        date_style = styles.add_style('CultivoDate', 4)
        date_style.font.name = 'Calibri'
        date_style.font.italic = True
        date_style.font.size = Pt(cfg["fonts"]["sizes"]["meta"])
        date_style.font.color.rgb = RGBColor.from_string(cfg["colors"]["secondary"][1:])
        
        # Cuerpo (texto negro)
        body_style = styles['Normal']
        body_style.font.name = 'Arial'
        body_style.font.size = Pt(cfg["fonts"]["sizes"]["body"])
        body_style.font.color.rgb = RGBColor(0, 0, 0)  # Negro puro
    
    def _add_cover_page(self, doc):
        # Logo
        if os.path.exists(PDF_CONFIG["branding"]["logo_path"]):
            doc.add_picture(PDF_CONFIG["branding"]["logo_path"], width=Inches(2))
        
        # Título
        title = doc.add_paragraph("Blog Cultivo Loco", style='Title')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_page_break()
    
    def _add_article(self, doc, article):
        # Título
        doc.add_heading(article['title'], level=1)
        
        # Fecha
        doc.add_paragraph(f"Publicado el {article['date']}", style='CultivoDate')
        
        # Contenido
        content = doc.add_paragraph(article['content'])
        content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        # Divisor
        doc.add_paragraph().add_run("―" * 50).color.rgb = RGBColor.from_string(PDF_CONFIG["colors"]["border"][1:])